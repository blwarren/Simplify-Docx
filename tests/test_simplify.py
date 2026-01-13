"""Tests for the simplify API."""

from __future__ import annotations

from docx import Document

from simplify_docx import simplify


def _build_document() -> Document:
    """Create a basic document for simplify tests."""
    doc = Document()
    doc.add_paragraph("Hello world")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    return doc


def test_simplify_returns_document_json() -> None:
    """Simplify returns a JSON-like mapping with a document type."""
    doc = _build_document()
    result = simplify(doc)

    assert isinstance(result, dict)
    assert result.get("TYPE") == "document"
    assert "VALUE" in result
    assert result["VALUE"][0].get("TYPE") == "body"
    body_value = result["VALUE"][0].get("VALUE", [])
    assert any(child.get("TYPE") == "paragraph" for child in body_value)


def test_simplify_can_disable_friendly_names() -> None:
    """Friendly names can be disabled via options."""
    doc = _build_document()
    result = simplify(doc, {"friendly-name": False})

    assert result.get("TYPE") == "CT_Document"
